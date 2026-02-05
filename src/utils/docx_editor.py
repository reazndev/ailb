from docx import Document
from docx.shared import RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
import re
from typing import List, Dict

def integrate_solution_to_docx(original_path: str, output_path: str, task_results: List[Dict[str, str]]):
    """
    Integrates AI solutions into the original DOCX document.
    Matches tasks to document locations and applies dark blue formatting.
    """
    print(f"DEBUG: Integrating {len(task_results)} tasks into {original_path}")
    
    try:
        if not os.path.exists(original_path):
            print(f"ERROR: Original file {original_path} not found.")
            return False
            
        doc = Document(original_path)
        DARK_BLUE = RGBColor(0x00, 0x00, 0x8B)
        
        # Helper to extract numbers from task title (e.g., "1. Aufgabe" -> 1)
        def get_task_number(title):
            match = re.search(r'(\d+)', title)
            return match.group(1) if match else None

        used_tasks = set()

        # 1. Try to fill tables first (common for "Aufgabe" | "Lösung" structure)
        for table in doc.tables:
            for row in table.rows:
                for i, cell in enumerate(row.cells):
                    cell_text = cell.text.strip().lower()
                    
                    # Look for tasks matching this cell
                    for idx, res in enumerate(task_results):
                        if idx in used_tasks: continue
                        
                        task_num = get_task_number(res['task'])
                        task_title_clean = res['task'].lower()
                        
                        # Match by number or title substring
                        is_match = False
                        if task_num and (f"aufgabe {task_num}" in cell_text or cell_text == task_num or cell_text == f"{task_num}."):
                            is_match = True
                        elif len(cell_text) > 3 and cell_text in task_title_clean:
                            is_match = True
                            
                        if is_match:
                            # If we matched a label cell, check if next cell is empty or has a placeholder
                            if i + 1 < len(row.cells):
                                target_cell = row.cells[i+1]
                                if not target_cell.text.strip() or "lösung" in target_cell.text.lower():
                                    # Fill it!
                                    target_cell.text = "" # Clear placeholder
                                    p = target_cell.paragraphs[0]
                                    run = p.add_run(res['content'])
                                    run.font.color.rgb = DARK_BLUE
                                    run.bold = True
                                    used_tasks.add(idx)
                                    print(f"DEBUG: Integrated Task {idx} into table cell.")
                                    break

        # 2. Try to match paragraphs
        # We iterate in reverse or carefully to insert after matches
        # For simplicity, we'll collect matches first
        for idx, res in enumerate(task_results):
            if idx in used_tasks: continue
            
            task_num = get_task_number(res['task'])
            task_title_clean = res['task'].lower()
            
            best_p_idx = -1
            
            for p_idx, p in enumerate(doc.paragraphs):
                p_text = p.text.strip().lower()
                if not p_text: continue
                
                match_found = False
                if task_num and (f"aufgabe {task_num}" in p_text or p_text == f"{task_num}." or p_text == task_num):
                    match_found = True
                elif len(p_text) > 10 and p_text in task_title_clean:
                    match_found = True
                
                if match_found:
                    best_p_idx = p_idx
                    # Don't break, find the LAST occurrence or keep this one?
                    # Usually first is fine for labels.
                    break
            
            if best_p_idx != -1:
                # Insert after this paragraph
                # python-docx doesn't have a direct 'insert_after', so we use internals or append
                # Safe way: add a paragraph and format it
                new_p = doc.paragraphs[best_p_idx].insert_paragraph_before("") # wait, this is before.
                # Actually, doc.paragraphs[best_p_idx]._element.addnext(...) is needed for after.
                # Simplest hack: Use the paragraph itself if it ends with "Lösung:" or similar
                # or just add at the end of matching logic.
                
                # Re-using the matched paragraph if it's a short label
                matched_p = doc.paragraphs[best_p_idx]
                if len(matched_p.text) < 50:
                    new_p = doc.add_paragraph()
                    matched_p._element.addnext(new_p._element)
                    
                    run = new_p.add_run("\n" + res['content'])
                    run.font.color.rgb = DARK_BLUE
                    used_tasks.add(idx)
                    print(f"DEBUG: Integrated Task {idx} after paragraph {best_p_idx}.")

        # 3. Fallback: Append remaining tasks at the end
        remaining = [res for i, res in enumerate(task_results) if i not in used_tasks]
        if remaining:
            print(f"DEBUG: Appending {len(remaining)} tasks at the end (no match found).")
            doc.add_page_break()
            header = doc.add_heading("Zusätzliche Ausarbeitungen", level=1)
            for res in remaining:
                p = doc.add_paragraph()
                p.add_run(f"{res['task']}\n").bold = True
                run = p.add_run(res['content'])
                run.font.color.rgb = DARK_BLUE
                doc.add_paragraph("") # Spacer

        # Save
        directory = os.path.dirname(output_path)
        if directory:
            os.makedirs(directory, exist_ok=True)
        doc.save(output_path)
        print(f"DEBUG: Successfully saved integrated DOCX to {output_path}")
        return True

    except Exception as e:
        print(f"Error integrating DOCX: {e}")
        import traceback
        traceback.print_exc()
        return False

# Backward compatibility wrapper
def append_solution_to_docx(original_path: str, output_path: str, task_results: List[Dict[str, str]]):
    return integrate_solution_to_docx(original_path, output_path, task_results)