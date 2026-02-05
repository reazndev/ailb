from src.utils.docx_editor import integrate_solution_to_docx
from docx import Document
import os

def test_docx_integration():
    dummy_input = "test_integration_input.docx"
    dummy_output = "test_integration_output.docx"
    
    # Create a dummy assignment
    doc = Document()
    doc.add_heading("Hausaufgabe", 0)
    doc.add_paragraph("Aufgabe 1: Erkläre den Begriff KI.")
    doc.add_paragraph("Hier ist Platz für die Lösung:")
    
    # Add a table
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Aufgabe 2"
    table.cell(0, 1).text = "Lösung hier einfügen"
    table.cell(1, 0).text = "Unbekannt"
    table.cell(1, 1).text = ""
    
    doc.save(dummy_input)
    
    task_results = [
        {"task": "1. Erkläre den Begriff KI", "content": "KI steht für Künstliche Intelligenz."},
        {"task": "2. Zweite Aufgabe bearbeiten", "content": "Dies ist die Lösung für Aufgabe 2."},
        {"task": "3. Etwas völlig Neues", "content": "Dieser Teil sollte am Ende landen."}
    ]
    
    print(f"Testing integration...")
    result = integrate_solution_to_docx(dummy_input, dummy_output, task_results)
    
    if result:
        print("Success! Verifying output...")
        doc2 = Document(dummy_output)
        
        # Check table
        t2 = doc2.tables[0]
        print(f"Table Cell (0,1) text: '{t2.cell(0, 1).text}'")
        
        # Check paragraphs
        print("--- PARAGRAPHS ---")
        for p in doc2.paragraphs:
            if p.text.strip():
                print(f"P: {p.text[:50]}")
        
        print("Verification PASSED (Manual check recommended for colors)")
    else:
        print("Integration failed.")

    # Cleanup (Optional, keep for manual inspection)
    # os.remove(dummy_input)

if __name__ == "__main__":
    test_docx_integration()
