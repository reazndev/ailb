from src.utils.docx_editor import append_solution_to_docx
from docx import Document
import os

def test_docx_append():
    # Create a dummy docx
    dummy_input = "test_input.docx"
    dummy_output = "test_output.docx"
    
    doc = Document()
    doc.add_paragraph("Original Content")
    doc.save(dummy_input)
    
    solution_text = "This is the generated solution."
    
    print(f"Appending to {dummy_input} -> {dummy_output}")
    result = append_solution_to_docx(dummy_input, dummy_output, solution_text)
    
    if result:
        print("Success!")
        # Verify
        doc2 = Document(dummy_output)
        full_text = "\n".join([p.text for p in doc2.paragraphs])
        print("--- CONTENT ---")
        print(full_text)
        print("---------------")
        
        if "AI Solution" in full_text and solution_text in full_text:
            print("Verification PASSED")
        else:
            print("Verification FAILED")
    else:
        print("Function returned False")

    # Cleanup
    if os.path.exists(dummy_input): os.remove(dummy_input)
    if os.path.exists(dummy_output): os.remove(dummy_output)

if __name__ == "__main__":
    test_docx_append()
